import argparse
import copy
import sys

from docx import Document


def parse_args():
    parser = argparse.ArgumentParser(
        description=(
            "Copy a table from a source DOCX into a template DOCX by replacing "
            "a placeholder paragraph."
        )
    )
    parser.add_argument("-s", "--source", required=True, help="Source DOCX path")
    parser.add_argument(
        "-t", "--template", required=True, help="Template DOCX path"
    )
    parser.add_argument("-o", "--output", required=True, help="Output DOCX path")
    parser.add_argument(
        "-p",
        "--placeholder",
        required=True,
        help="Placeholder text to find and replace",
    )
    parser.add_argument(
        "--table-index",
        type=int,
        default=1,
        help="1-based table index from source DOCX (default: 1)",
    )
    return parser.parse_args()


def main():
    args = parse_args()

    if args.table_index < 1:
        print("ERROR: --table-index must be >= 1", file=sys.stderr)
        return 2

    src = Document(args.source)
    if len(src.tables) < args.table_index:
        print(
            (
                "ERROR: source document has "
                f"{len(src.tables)} table(s), but --table-index={args.table_index}"
            ),
            file=sys.stderr,
        )
        return 2

    table_el = copy.deepcopy(src.tables[args.table_index - 1]._element)
    tpl = Document(args.template)

    target_para = None
    for para in tpl.paragraphs:
        if args.placeholder in para.text:
            target_para = para
            break

    if target_para is None:
        print(
            f"ERROR: placeholder not found: {args.placeholder}",
            file=sys.stderr,
        )
        return 2

    target_el = target_para._element
    target_el.addprevious(table_el)
    target_el.getparent().remove(target_el)

    tpl.save(args.output)
    print(f"Done: table inserted into {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
