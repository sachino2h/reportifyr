import argparse
import copy
import os
import re
import sys
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

import helper


INLINE_PATTERN = re.compile(r"\{\{([^{}]+)\}\}|<<text:([^>]+)>>")
IMAGE_PATTERN = re.compile(r"<<image:([^>]+)>>")
TABLE_PATTERN = re.compile(r"<<table:([^>]+)>>")
GENERIC_BLOCK_PATTERN = re.compile(r"<<([^<>:|]+)>>")
EXPANDED_BLOCK_PATTERN = re.compile(r"<<([^<>]+\|[^<>]+)>>")
VALID_IMAGE_EXT = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff"}
VALID_TABLE_EXT = {".csv", ".rds", ".docx"}
MISSING_PATTERN = re.compile(r"^MISSING_VALUE\(.+\)$", re.IGNORECASE)


def normalize_key(value):
    return re.sub(r"\s+", " ", str(value or "").strip()).lower()


def canonical_key(value):
    s = normalize_key(value)
    s = s.strip("\"'")
    s = re.sub(r"[^a-z0-9]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def slugify(value):
    s = re.sub(r"[^a-zA-Z0-9]+", "_", str(value or "").strip().lower())
    s = s.strip("_")
    return s or "block"


def marker_text(name, block_id):
    return f"RPFY_{name}:{block_id}"


def marker_text_with_value(name, block_id, value):
    return f"RPFY_{name}:{block_id}:{value}"


def add_hidden_marker(paragraph, text):
    run = paragraph.add_run(text)
    run.font.hidden = True


def is_missing_token(text):
    return bool(MISSING_PATTERN.match(str(text or "").strip()))


def insert_run_after(paragraph, anchor_run, text, hidden=False, style_from_run=None):
    new_run = paragraph.add_run(text)

    src = style_from_run
    if src is not None and src._r.rPr is not None:
        if new_run._r.rPr is not None:
            new_run._r.remove(new_run._r.rPr)
        new_run._r.insert(0, copy.deepcopy(src._r.rPr))

    if hidden:
        new_run.font.hidden = True

    anchor_el = anchor_run._r
    new_el = new_run._r
    new_el.getparent().remove(new_el)
    anchor_el.addnext(new_el)
    return new_run


def add_title_run(paragraph, title_text):
    run = paragraph.add_run(title_text)
    run.bold = True
    run.font.size = Pt(12)
    if is_missing_token(title_text):
        run.font.color.rgb = RGBColor(255, 0, 0)
    return run


def add_footnote_run(paragraph, footnote_text):
    run = paragraph.add_run(footnote_text)
    run.font.size = Pt(10)
    if is_missing_token(footnote_text):
        run.font.color.rgb = RGBColor(255, 0, 0)
    return run


def add_missing_run(paragraph, text):
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(255, 0, 0)
    return run


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def paragraph_replace_inline(paragraph, inline_map, inline_seen):
    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(r.text for r in runs)
    if not full_text:
        return

    matches = list(INLINE_PATTERN.finditer(full_text))
    if not matches:
        return

    bounds = []
    pos = 0
    for run in runs:
        start = pos
        end = pos + len(run.text)
        bounds.append((start, end))
        pos = end

    def find_run_index(char_pos):
        for i, (start, end) in enumerate(bounds):
            if start <= char_pos < end:
                return i
        if bounds and char_pos == bounds[-1][1]:
            return len(bounds) - 1
        return None

    for m in reversed(matches):
        key = (m.group(1) or m.group(2) or "").strip()
        entry = (
            inline_map.get(key)
            or inline_map.get(normalize_key(key))
            or inline_map.get(canonical_key(key))
        )
        if isinstance(entry, dict):
            repl = str(entry.get("value", ""))
        elif entry is None:
            repl = ""
        else:
            repl = str(entry)

        start_pos, end_pos = m.start(), m.end()
        start_run = find_run_index(start_pos)
        end_run = find_run_index(end_pos - 1) if end_pos > start_pos else start_run
        if start_run is None or end_run is None:
            continue

        start_text = runs[start_run].text
        start_bound = bounds[start_run][0]
        start_offset = start_pos - start_bound

        end_text = runs[end_run].text
        end_bound = bounds[end_run][0]
        end_offset = end_pos - end_bound

        preferred_id = None
        if isinstance(entry, dict):
            preferred_id = str(entry.get("id", "")).strip() or None
        base = preferred_id or f"inl_{slugify(key)}"
        count = inline_seen.get(base, 0) + 1
        inline_seen[base] = count
        inline_id = base if count == 1 else f"{base}_{count}"

        prefix = start_text[:start_offset]
        suffix = end_text[end_offset:]

        # Remove matched text from original runs, then insert marker/value runs
        # directly after the prefix run to keep strict order:
        # START -> visible value -> END -> KEY.
        runs[start_run].text = prefix
        if start_run != end_run:
            for i in range(start_run + 1, end_run):
                runs[i].text = ""
            runs[end_run].text = suffix

        anchor = runs[start_run]
        h_start = insert_run_after(
            paragraph,
            anchor,
            marker_text("INLINE_START", inline_id),
            hidden=True
        )
        v_run = insert_run_after(
            paragraph,
            h_start,
            repl,
            hidden=False,
            style_from_run=runs[start_run]
        )
        if is_missing_token(repl):
            v_run.font.color.rgb = RGBColor(255, 0, 0)
        h_end = insert_run_after(
            paragraph,
            v_run,
            marker_text("INLINE_END", inline_id),
            hidden=True
        )
        h_key = insert_run_after(
            paragraph,
            h_end,
            marker_text_with_value("INLINE_KEY", inline_id, key),
            hidden=True
        )

        if start_run == end_run and suffix:
            insert_run_after(
                paragraph,
                h_key,
                suffix,
                hidden=False,
                style_from_run=runs[start_run]
            )
            runs[start_run].text = prefix

        # Recompute state for remaining matches.
        runs = paragraph.runs
        full_text = "".join(r.text for r in runs)
        bounds = []
        pos = 0
        for run in runs:
            start = pos
            end = pos + len(run.text)
            bounds.append((start, end))
            pos = end


def paragraph_replace_inline_text_nodes_only(paragraph, inline_map):
    # Fallback pass for field/hyperlink areas where python-docx run handling can
    # miss visible text nodes (for example TOC/TOF/pagination result text).
    text_nodes = paragraph._element.xpath(".//w:t")
    if not text_nodes:
        return

    parts = [(node.text or "") for node in text_nodes]
    original = "".join(parts)
    if not original:
        return

    def repl(match):
        key = (match.group(1) or match.group(2) or "").strip()
        entry = (
            inline_map.get(key)
            or inline_map.get(normalize_key(key))
            or inline_map.get(canonical_key(key))
        )
        if isinstance(entry, dict):
            value = str(entry.get("value", ""))
        elif entry is None:
            value = ""
        else:
            value = str(entry)
        return value

    updated = INLINE_PATTERN.sub(repl, original)
    if updated == original:
        return

    lengths = [len(p) for p in parts]
    pos = 0
    for i, node in enumerate(text_nodes):
        if i < len(text_nodes) - 1:
            take = lengths[i]
            node.text = updated[pos : pos + take]
            pos += take
        else:
            node.text = updated[pos:]


def parse_expanded_block(raw_content):
    # Expected: Name:...|Title:...|Footnote:...|Files:...|Type:...
    parts = [p.strip() for p in raw_content.split("|")]
    data = {}
    for part in parts:
        if ":" not in part:
            continue
        k, v = part.split(":", 1)
        data[normalize_key(k)] = v.strip()
    if not data:
        return None
    if "name" not in data:
        return None
    return data


def parse_block_tag(paragraph_text):
    m = IMAGE_PATTERN.search(paragraph_text)
    if m:
        return {
            "key": m.group(1).strip(),
            "type": "image",
            "raw_tag": m.group(0),
            "title": "",
            "footnote": "",
            "files": [],
        }

    tm = TABLE_PATTERN.search(paragraph_text)
    if tm:
        return {
            "key": tm.group(1).strip(),
            "type": "table",
            "raw_tag": tm.group(0),
            "title": "",
            "footnote": "",
            "files": [],
        }

    em = EXPANDED_BLOCK_PATTERN.search(paragraph_text)
    if em:
        expanded = parse_expanded_block(em.group(1))
        if expanded and expanded.get("type", "").lower() in ("image", "table"):
            files = []
            files_raw = expanded.get("files", "")
            if files_raw:
                files = [f.strip() for f in files_raw.split(",") if f.strip()]
            return {
                "key": expanded.get("name", "").strip(),
                "type": expanded.get("type", "").lower(),
                "raw_tag": em.group(0),
                "title": expanded.get("title", ""),
                "footnote": expanded.get("footnote", ""),
                "files": files,
            }

    gm = GENERIC_BLOCK_PATTERN.search(paragraph_text)
    if gm:
        key = gm.group(1).strip()
        return {
            "key": key,
            "type": "",
            "raw_tag": gm.group(0),
            "title": "",
            "footnote": "",
            "files": [],
        }
    return None


def block_lookup(blocks_map, key):
    return blocks_map.get(key) or blocks_map.get(normalize_key(key)) or {}


def coerce_files(value):
    if value is None:
        return []
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    if isinstance(value, str):
        return [v.strip() for v in value.split(",") if v.strip()]
    return []


def infer_block_type_from_files(files):
    has_image = False
    has_table = False
    for file_name in files:
        parsed = parse_file_entry(file_name)
        base = os.path.basename(parsed["file"])
        ext = os.path.splitext(base)[1].lower()
        if ext in VALID_IMAGE_EXT:
            has_image = True
        if ext in VALID_TABLE_EXT:
            has_table = True
    if has_table and not has_image:
        return "table"
    if has_image and not has_table:
        return "image"
    return ""


def parse_file_entry(value):
    # Supports entries like "plot.png<width: 5, height: 3>"
    text = str(value or "").strip()
    if not text:
        return {"file": "", "args": {}}

    m = re.match(r"^(.*?)(?:<(.*)>)?$", text)
    file_name = (m.group(1) if m else text).strip()
    args = {}
    args_txt = (m.group(2) if m else None)
    if args_txt:
        for part in args_txt.split(","):
            if ":" not in part:
                continue
            key, val = part.split(":", 1)
            args[key.strip().lower()] = val.strip()
    return {"file": file_name, "args": args}


def unique_block_id(preferred_id, key, used_ids):
    base = preferred_id or f"blk_{slugify(key)}"
    if base not in used_ids:
        used_ids.add(base)
        return base
    i = 2
    while f"{base}_{i}" in used_ids:
        i += 1
    final_id = f"{base}_{i}"
    used_ids.add(final_id)
    return final_id


def resolve_images(files, assets_dir):
    resolved = []
    for file_name in files:
        parsed = parse_file_entry(file_name)
        base = os.path.basename(parsed["file"])
        path = os.path.join(assets_dir, base)
        ext = os.path.splitext(base)[1].lower()
        if os.path.exists(path) and ext in VALID_IMAGE_EXT:
            resolved.append({"path": path, "args": parsed["args"]})
    return resolved


def resolve_table_placeholder(files, key=None, tables_dir=None):
    for file_name in files:
        parsed = parse_file_entry(file_name)
        base = os.path.basename(parsed["file"])
        ext = os.path.splitext(base)[1].lower()
        if ext in VALID_TABLE_EXT:
            return base

    if key:
        base = os.path.basename(str(key).strip())
        ext = os.path.splitext(base)[1].lower()
        if ext in VALID_TABLE_EXT:
            return base

    if key and tables_dir and os.path.isdir(tables_dir):
        desired = normalize_key(os.path.splitext(os.path.basename(str(key)))[0])
        candidates = []
        for name in os.listdir(tables_dir):
            ext = os.path.splitext(name)[1].lower()
            if ext not in VALID_TABLE_EXT:
                continue
            stem = normalize_key(os.path.splitext(name)[0])
            if stem == desired:
                candidates.append(name)
        if candidates:
            candidates.sort(key=lambda x: (0 if x.lower().endswith(".docx") else 1, x.lower()))
            return candidates[0]
    return ""


def add_picture_with_config(run, image_path, image_args, config):
    use_embedded_size = bool(config.get("use_embedded_size", True))
    use_artifact_size = bool(config.get("use_artifact_size", False))

    has_width = "width" in image_args and str(image_args["width"]).strip() != ""
    has_height = "height" in image_args and str(image_args["height"]).strip() != ""

    if use_embedded_size and (has_width or has_height):
        run.add_picture(
            image_path,
            width=Inches(float(image_args["width"])) if has_width else None,
            height=Inches(float(image_args["height"])) if has_height else None,
        )
        return

    if use_artifact_size:
        run.add_picture(image_path)
        return

    default_width = float(config.get("default_fig_width", 6))
    if has_width or has_height:
        run.add_picture(
            image_path,
            width=Inches(float(image_args["width"])) if has_width else None,
            height=Inches(float(image_args["height"])) if has_height else None,
        )
    else:
        run.add_picture(image_path, width=Inches(default_width))


def resolve_alignment(config):
    value = str(config.get("fig_alignment", "center")).strip().lower()
    if value == "left":
        return WD_ALIGN_PARAGRAPH.LEFT
    if value == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.CENTER


def remove_tag_from_paragraph(paragraph, tag):
    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(r.text for r in runs)
    idx = full_text.find(tag)
    if idx < 0:
        return

    start_pos = idx
    end_pos = idx + len(tag)

    bounds = []
    pos = 0
    for run in runs:
        start = pos
        end = pos + len(run.text)
        bounds.append((start, end))
        pos = end

    def find_run_index(char_pos):
        for i, (start, end) in enumerate(bounds):
            if start <= char_pos < end:
                return i
        return None

    start_run = find_run_index(start_pos)
    end_run = find_run_index(end_pos - 1) if end_pos > start_pos else start_run
    if start_run is None or end_run is None:
        return

    start_text = runs[start_run].text
    start_bound = bounds[start_run][0]
    start_offset = start_pos - start_bound

    end_text = runs[end_run].text
    end_bound = bounds[end_run][0]
    end_offset = end_pos - end_bound

    if start_run == end_run:
        runs[start_run].text = start_text[:start_offset] + start_text[end_offset:]
    else:
        runs[start_run].text = start_text[:start_offset]
        for i in range(start_run + 1, end_run):
            runs[i].text = ""
        runs[end_run].text = end_text[end_offset:]


def restore_preserved_docx_parts(docx_in, docx_out):
    # Preserve document styling/layout parts exactly from input.
    preserve_prefixes = (
        "word/header",
        "word/footer",
        "word/styles",
        "word/settings",
        "word/theme/",
        "word/fontTable",
        "word/numbering",
        "docProps/",
    )
    preserve_exact = {"_rels/.rels"}

    with zipfile.ZipFile(docx_in, "r") as zin:
        source_parts = {name: zin.read(name) for name in zin.namelist()}
    with zipfile.ZipFile(docx_out, "r") as zout_in:
        out_parts = {name: zout_in.read(name) for name in zout_in.namelist()}

    for name, payload in source_parts.items():
        should_copy = name in preserve_exact or name.startswith(preserve_prefixes)
        if should_copy:
            out_parts[name] = payload

    tmp_path = f"{docx_out}.tmp"
    with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as ztmp:
        for name, payload in out_parts.items():
            ztmp.writestr(name, payload)
    os.replace(tmp_path, docx_out)


def process_doc(docx_in, docx_out, yaml_in, assets_dir, tables_dir=None, config_yaml=None):
    doc = Document(docx_in)
    yml = helper.load_yaml(yaml_in) or {}
    config = helper.load_yaml(config_yaml) if config_yaml else {}
    inline_raw = (
        yml.get("inline")
        or yml.get("inlines")
        or yml.get("inline_block")
        or {}
    )
    blocks_raw = yml.get("blocks") or {}

    inline_map = dict(inline_raw)
    for k, v in list(inline_raw.items()):
        inline_map[normalize_key(k)] = v
        inline_map[canonical_key(k)] = v

    blocks_map = dict(blocks_raw)
    for k, v in list(blocks_raw.items()):
        blocks_map[normalize_key(k)] = v

    used_ids = set()

    # Inline text replacement first.
    inline_seen = {}
    for paragraph in list(iter_all_paragraphs(doc)):
        paragraph_replace_inline(paragraph, inline_map, inline_seen)
        paragraph_replace_inline_text_nodes_only(paragraph, inline_map)

    # Block rendering next.
    for paragraph in list(iter_all_paragraphs(doc)):
        tag = parse_block_tag(paragraph.text)
        if not tag:
            continue

        block_yaml = block_lookup(blocks_map, tag["key"])
        block_type = (tag["type"] or block_yaml.get("type") or "").strip().lower()
        block_id = unique_block_id(block_yaml.get("id"), tag["key"], used_ids)
        title = (tag.get("title") or block_yaml.get("title") or "").strip()
        footnote = (tag.get("footnote") or block_yaml.get("footnote") or "").strip()
        files = tag.get("files") or coerce_files(block_yaml.get("files"))
        inferred_type = infer_block_type_from_files(files)
        if inferred_type:
            block_type = inferred_type
        if block_type not in ("image", "table"):
            continue

        remove_tag_from_paragraph(paragraph, tag["raw_tag"])

        add_hidden_marker(paragraph, marker_text("BLOCK_START", block_id))
        add_hidden_marker(
            paragraph,
            marker_text_with_value("BLOCK_KEY", block_id, slugify(tag["key"])),
        )
        add_hidden_marker(
            paragraph,
            marker_text_with_value("BLOCK_TYPE", block_id, block_type),
        )
        add_hidden_marker(paragraph, marker_text("TITLE_START", block_id))
        if title:
            add_title_run(paragraph, title)
        add_hidden_marker(paragraph, marker_text("TITLE_END", block_id))

        anchor = paragraph
        if block_type == "image":
            resolved_images = resolve_images(files, assets_dir)
            add_hidden_marker(anchor, marker_text("IMAGE_START", block_id))
            for image_info in resolved_images:
                image_para = insert_paragraph_after(anchor)
                image_para.alignment = resolve_alignment(config)
                run = image_para.add_run()
                add_picture_with_config(
                    run=run,
                    image_path=image_info["path"],
                    image_args=image_info["args"],
                    config=config,
                )
                add_hidden_marker(
                    image_para,
                    marker_text_with_value(
                        "IMAGE_FILE",
                        block_id,
                        os.path.basename(image_info["path"]),
                    ),
                )
                anchor = image_para
            if len(resolved_images) == 0:
                missing_files = [f for f in files if is_missing_token(f)]
                if missing_files:
                    image_para = insert_paragraph_after(anchor)
                    image_para.alignment = resolve_alignment(config)
                    add_missing_run(image_para, ", ".join(missing_files))
                    anchor = image_para
            add_hidden_marker(anchor, marker_text("IMAGE_END", block_id))
        else:
            # Extended table tags are matched as <<table:name>> (or expanded Type:table).
            # We intentionally create a temporary {rpfy}:file.ext placeholder so the
            # existing R add_tables() pipeline can be reused without duplicating table
            # insertion logic. The temporary placeholder is removed after add_tables().
            add_hidden_marker(anchor, marker_text("TABLE_START", block_id))
            table_file_for_marker = ""
            placeholder_file = resolve_table_placeholder(files, key=tag["key"], tables_dir=tables_dir)
            if placeholder_file:
                placeholder = f"{{rpfy}}:{placeholder_file}"
                table_para = insert_paragraph_after(anchor)
                table_para.style = paragraph.style
                table_para.add_run(placeholder)
                table_file_for_marker = placeholder_file
                # Keep anchor at placeholder so footnote is placed after rendered table.
                anchor = table_para
            else:
                missing_files = [f for f in files if is_missing_token(f)]
                if missing_files:
                    miss_para = insert_paragraph_after(anchor)
                    miss_para.style = paragraph.style
                    add_missing_run(miss_para, ", ".join(missing_files))
                    anchor = miss_para

        foot_para = insert_paragraph_after(anchor)
        add_hidden_marker(foot_para, marker_text("FOOTNOTE_START", block_id))
        if footnote:
            foot_para.style = paragraph.style
            add_footnote_run(foot_para, footnote)
        add_hidden_marker(foot_para, marker_text("FOOTNOTE_END", block_id))
        if block_type == "table":
            if table_file_for_marker:
                add_hidden_marker(
                    foot_para,
                    marker_text_with_value("TABLE_FILE", block_id, table_file_for_marker),
                )
            add_hidden_marker(foot_para, marker_text("TABLE_END", block_id))
        add_hidden_marker(foot_para, marker_text("BLOCK_END", block_id))

    doc.save(docx_out)
    restore_preserved_docx_parts(docx_in, docx_out)
    print(f"Processed file saved at '{docx_out}'.")


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        yield from iter_table_paragraphs(tbl)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def main():
    parser = argparse.ArgumentParser(description="Apply YAML inline and image blocks to DOCX")
    parser.add_argument("-i", "--input", required=True, help="Input DOCX path")
    parser.add_argument("-o", "--output", required=True, help="Output DOCX path")
    parser.add_argument("-y", "--yaml", required=True, help="YAML path")
    parser.add_argument("-a", "--assets_dir", required=True, help="Assets directory")
    parser.add_argument("-t", "--tables_dir", required=False, default=None, help="Tables directory")
    parser.add_argument("-c", "--config", required=False, default=None, help="Config yaml path")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        raise FileNotFoundError(f"Input DOCX not found: {args.input}")
    if not os.path.exists(args.yaml):
        raise FileNotFoundError(f"YAML not found: {args.yaml}")
    if not os.path.isdir(args.assets_dir):
        raise NotADirectoryError(f"assets_dir is not a directory: {args.assets_dir}")

    process_doc(
        args.input,
        args.output,
        args.yaml,
        args.assets_dir,
        tables_dir=args.tables_dir,
        config_yaml=args.config,
    )


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(str(e), file=sys.stderr)
        sys.exit(1)
